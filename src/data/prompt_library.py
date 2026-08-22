"""
提示词模板库
扫描项目根目录的 Prompt和skills库/ 文件夹，把 docx 指令包接入为提示词模板
"""

import os
import re
from dataclasses import dataclass
from typing import Dict, List, Optional

_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DEFAULT_PROMPT_DIR = os.path.join(_PROJECT_ROOT, "Prompt和skills库")


@dataclass
class PromptTemplate:
    """一个提示词模板"""
    name: str           # 模板名（去掉后缀的文件名）
    pack: str           # 所属指令包（父目录名）
    file_path: str      # docx 路径


class PromptLibrary:
    """提示词模板库"""

    def __init__(self, prompt_dir: Optional[str] = None):
        self.prompt_dir = prompt_dir or DEFAULT_PROMPT_DIR
        self._templates: List[PromptTemplate] = []
        self._text_cache: Dict[str, str] = {}
        self.refresh()

    def refresh(self):
        """重新扫描（只索引 docx/doc 指令文档，跳过视频与图片）"""
        self._templates = []
        if not os.path.isdir(self.prompt_dir):
            return

        for root, _dirs, files in os.walk(self.prompt_dir):
            for fname in sorted(files):
                if not fname.lower().endswith((".docx",)):
                    continue
                pack = os.path.basename(root)
                if pack == os.path.basename(self.prompt_dir):
                    pack = "根目录"
                name = os.path.splitext(fname)[0]
                self._templates.append(PromptTemplate(
                    name=name,
                    pack=pack,
                    file_path=os.path.join(root, fname),
                ))

    def list_packs(self) -> List[str]:
        """所有指令包名"""
        return sorted({t.pack for t in self._templates})

    def list_templates(self, pack: Optional[str] = None) -> List[PromptTemplate]:
        """模板列表，可按指令包过滤"""
        if pack:
            return [t for t in self._templates if t.pack == pack]
        return list(self._templates)

    def search(self, keyword: str) -> List[PromptTemplate]:
        """按名称模糊搜索"""
        keyword = keyword.lower()
        return [t for t in self._templates
                if keyword in t.name.lower() or keyword in t.pack.lower()]

    def count(self) -> int:
        return len(self._templates)

    def read_text(self, template: PromptTemplate) -> str:
        """提取 docx 文本（带缓存）"""
        if template.file_path in self._text_cache:
            return self._text_cache[template.file_path]

        text = self._extract_docx_text(template.file_path)
        self._text_cache[template.file_path] = text
        return text

    @staticmethod
    def _extract_docx_text(file_path: str) -> str:
        """从 docx 提取纯文本"""
        try:
            import docx
        except ImportError:
            return "(缺少 python-docx 依赖，请 pip install python-docx)"

        try:
            document = docx.Document(file_path)
        except Exception as e:
            return f"(文档解析失败：{e})"

        paragraphs = []
        for para in document.paragraphs:
            line = para.text.strip()
            if line:
                paragraphs.append(line)

        # 表格内容也提取（部分指令写在表格里）
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    line = cell.text.strip()
                    if line and line not in paragraphs:
                        paragraphs.append(line)

        return "\n".join(paragraphs) if paragraphs else "(文档为空)"


_library: Optional[PromptLibrary] = None


def get_prompt_library() -> PromptLibrary:
    """获取提示词库单例"""
    global _library
    if _library is None:
        _library = PromptLibrary()
    return _library
