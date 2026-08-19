"""
导出模块
支持将小说导出为TXT、Markdown、Word格式
"""

import os
from datetime import datetime
from typing import Optional


class NovelExporter:
    """小说导出器"""
    
    def __init__(self, output_dir: str = "data/exports"):
        """
        初始化导出器
        
        Args:
            output_dir: 导出文件存储目录
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def export_txt(self, title: str, chapters: list[dict], 
                   filename: Optional[str] = None) -> str:
        """
        导出为TXT格式
        
        Args:
            title: 小说标题
            chapters: 章节列表
            filename: 文件名（可选）
        
        Returns:
            导出文件路径
        """
        if not filename:
            filename = f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        
        filepath = os.path.join(self.output_dir, filename)
        
        with open(filepath, "w", encoding="utf-8") as f:
            # 写入标题
            f.write(f"{title}\n")
            f.write("=" * 50 + "\n\n")
            
            # 写入章节
            for chapter in chapters:
                chapter_num = chapter.get("number", 0)
                chapter_title = chapter.get("title", f"第{chapter_num}章")
                content = chapter.get("content", "")
                
                f.write(f"\n{'=' * 30}\n")
                f.write(f"{chapter_title}\n")
                f.write(f"{'=' * 30}\n\n")
                f.write(content)
                f.write("\n\n")
            
            # 写入结尾
            f.write("\n" + "=" * 50 + "\n")
            f.write("— 全文完 —\n")
        
        print(f"[导出] TXT文件已保存: {filepath}")
        return filepath
    
    def export_markdown(self, title: str, chapters: list[dict], 
                        metadata: dict = None, filename: Optional[str] = None) -> str:
        """
        导出为Markdown格式
        
        Args:
            title: 小说标题
            chapters: 章节列表
            metadata: 元数据
            filename: 文件名（可选）
        
        Returns:
            导出文件路径
        """
        if not filename:
            filename = f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        
        filepath = os.path.join(self.output_dir, filename)
        
        with open(filepath, "w", encoding="utf-8") as f:
            # 写入YAML头部
            f.write("---\n")
            f.write(f"title: {title}\n")
            if metadata:
                for key, value in metadata.items():
                    f.write(f"{key}: {value}\n")
            f.write(f"export_date: {datetime.now().isoformat()}\n")
            f.write("---\n\n")
            
            # 写入标题
            f.write(f"# {title}\n\n")
            
            # 写入目录
            f.write("## 目录\n\n")
            for chapter in chapters:
                chapter_num = chapter.get("number", 0)
                chapter_title = chapter.get("title", f"第{chapter_num}章")
                f.write(f"- [{chapter_title}](#{chapter_title})\n")
            f.write("\n---\n\n")
            
            # 写入章节
            for chapter in chapters:
                chapter_num = chapter.get("number", 0)
                chapter_title = chapter.get("title", f"第{chapter_num}章")
                content = chapter.get("content", "")
                summary = chapter.get("summary", "")
                
                f.write(f"## {chapter_title}\n\n")
                
                if summary:
                    f.write(f"> {summary}\n\n")
                
                f.write(content)
                f.write("\n\n---\n\n")
            
            # 写入结尾
            f.write("\n# — 全文完 —\n")
        
        print(f"[导出] Markdown文件已保存: {filepath}")
        return filepath
    
    def export_word(self, title: str, chapters: list[dict], 
                    metadata: dict = None, filename: Optional[str] = None) -> str:
        """
        导出为Word格式（需要python-docx）
        
        Args:
            title: 小说标题
            chapters: 章节列表
            metadata: 元数据
            filename: 文件名（可选）
        
        Returns:
            导出文件路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("[导出] 请安装python-docx: pip install python-docx")
            return ""
        
        if not filename:
            filename = f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        filepath = os.path.join(self.output_dir, filename)
        
        # 创建文档
        doc = Document()
        
        # 设置标题
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加元数据
        if metadata:
            doc.add_paragraph("")
            for key, value in metadata.items():
                p = doc.add_paragraph(f"{key}: {value}")
                p.style.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_page_break()
        
        # 添加目录标题
        doc.add_heading("目录", level=1)
        for chapter in chapters:
            chapter_num = chapter.get("number", 0)
            chapter_title = chapter.get("title", f"第{chapter_num}章")
            doc.add_paragraph(chapter_title, style="List Number")
        
        doc.add_page_break()
        
        # 添加章节内容
        for chapter in chapters:
            chapter_num = chapter.get("number", 0)
            chapter_title = chapter.get("title", f"第{chapter_num}章")
            content = chapter.get("content", "")
            
            # 添加章节标题
            doc.add_heading(chapter_title, level=1)
            
            # 添加内容
            paragraphs = content.split("\n")
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text)
            
            doc.add_page_break()
        
        # 保存文档
        doc.save(filepath)
        
        print(f"[导出] Word文件已保存: {filepath}")
        return filepath
    
    def export_pdf(self, title: str, chapters: list[dict],
                   metadata: dict = None, filename: Optional[str] = None) -> str:
        """
        导出为PDF（先生成docx，再用Word COM转换；无Word时降级为docx）
        
        Args:
            title: 小说标题
            chapters: 章节列表
            metadata: 元数据
            filename: 文件名（可选）
        
        Returns:
            导出文件路径（PDF或降级的DOCX）
        """
        if not filename:
            filename = f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        # 先导出 docx
        docx_path = self.export_word(title, chapters, metadata, filename=f"{filename}.docx")
        if not docx_path:
            return ""
        
        pdf_path = os.path.join(self.output_dir, f"{filename}.pdf")
        
        # 尝试用 Word COM 转换
        try:
            import win32com.client
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(os.path.abspath(docx_path))
                doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
                doc.Close(False)
                print(f"[导出] PDF文件已保存: {pdf_path}")
                return pdf_path
            finally:
                word.Quit()
        except Exception as e:
            print(f"[导出] 未找到Word/转换失败({e})，保留DOCX: {docx_path}")
            return docx_path
    
    def get_export_path(self, title: str, format: str) -> str:
        """获取导出文件路径"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        extensions = {"txt": ".txt", "markdown": ".md", "word": ".docx"}
        ext = extensions.get(format, ".txt")
        return os.path.join(self.output_dir, f"{title}_{timestamp}{ext}")


# 全局导出器实例
_exporter: Optional[NovelExporter] = None


def get_exporter() -> NovelExporter:
    """获取导出器单例"""
    global _exporter
    if _exporter is None:
        _exporter = NovelExporter()
    return _exporter